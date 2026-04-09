"""
Word 文档生成工具

功能：
- 生成带标题/段落/表格的 .docx 文档
- AI 根据指令填充内容
- 支持通过邮件发送

用法（指令通道）：
  Aegis: 帮我写一份会议纪要，内容是...
  Aegis: 生成一份项目进展表，项目A完成80%，项目B完成40%
  Aegis: 写一份邮件草稿并发给我
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import config

OUTPUT_DIR = config.DATA_DIR / "documents"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _safe_filename(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*\r\n\t]', '_', name).strip('_')[:60]


def create_word_doc(
    title: str,
    content: str,
    table_data: list[list[str]] = None,
    table_headers: list[str] = None,
    filename: str = None,
) -> Path:
    """
    生成 Word 文档。

    title: 文档标题
    content: 正文内容（Markdown 风格，## 变二级标题，- 变列表）
    table_data: 表格数据，list of rows，每行是字符串列表
    table_headers: 表格表头（可选）
    filename: 输出文件名（不含扩展名），默认按标题+时间生成

    返回生成的文件 Path。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── 页面边距 ─────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

    # ── 标题 ─────────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(18)

    # 副标题（日期）
    date_para = doc.add_paragraph(datetime.now().strftime("%Y年%m月%d日"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    date_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # ── 正文解析（简单 Markdown → Word）────────────────────────
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            doc.add_paragraph(text, style='List Number')
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
        else:
            # 处理行内加粗 **text**
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.+?)\*\*', stripped)
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.bold = (i % 2 == 1)

    # ── 表格 ─────────────────────────────────────────────────────
    if table_data:
        doc.add_paragraph()
        headers = table_headers or []
        rows_to_add = ([headers] + table_data) if headers else table_data

        if rows_to_add:
            col_count = max(len(r) for r in rows_to_add)
            table = doc.add_table(rows=len(rows_to_add), cols=col_count)
            table.style = 'Table Grid'

            for row_idx, row_data in enumerate(rows_to_add):
                row = table.rows[row_idx]
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < col_count:
                        cell = row.cells[col_idx]
                        cell.text = str(cell_text)
                        # 表头行加粗
                        if row_idx == 0 and headers:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                            # 表头背景色（浅蓝）
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), 'DCE6F1')
                            shd.set(qn('w:val'), 'clear')
                            tcPr.append(shd)

    # ── 保存 ─────────────────────────────────────────────────────
    if not filename:
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{_safe_filename(title)}_{ts}"

    out_path = OUTPUT_DIR / f"{filename}.docx"
    doc.save(str(out_path))
    return out_path


def ai_generate_doc(instruction: str) -> tuple[Path, str]:
    """
    让 AI 根据指令生成 Word 文档内容，返回 (文件路径, 描述)。

    AI 负责决定：标题、正文结构、是否需要表格及表格内容。
    """
    from ai import client as ai
    import json

    prompt = f"""用户指令：{instruction}

请生成一份 Word 文档的结构化内容，以 JSON 格式输出：
{{
  "title": "文档标题",
  "content": "正文内容（支持 ## 二级标题、- 列表、**加粗**，使用\\n换行）",
  "has_table": true或false,
  "table_headers": ["列1", "列2", "列3"],
  "table_data": [["值1", "值2", "值3"], ["值4", "值5", "值6"]],
  "description": "一句话说明生成了什么文档"
}}

要求：
- 正文要有实质内容，不要生成空白模板
- 根据指令判断是否需要表格，如不需要则 has_table: false，table_headers/table_data 为空
- 内容要专业、格式清晰
只输出 JSON。"""

    try:
        raw = ai.chat(
            messages=[{"role": "user", "content": prompt}],
            system_prompt="你是Aegis，用户的AI助理，擅长生成专业文档。",
            temperature=0.4,
        )
        raw = raw.strip().strip("```json").strip("```").strip()
        data = json.loads(raw)
    except Exception as e:
        # 兜底：把指令本身作为正文
        data = {
            "title": "文档",
            "content": instruction,
            "has_table": False,
            "table_headers": [],
            "table_data": [],
            "description": "根据指令生成的文档",
        }

    path = create_word_doc(
        title=data.get("title", "文档"),
        content=data.get("content", ""),
        table_headers=data.get("table_headers") or None,
        table_data=data.get("table_data") or None,
    )
    return path, data.get("description", "文档已生成")
